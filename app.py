# Essential imports for document handling, embedding, semantic search, and frontend
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.docstore.document import Document
from langchain.text_splitter import CharacterTextSplitter

# File handling and format-specific libraries
import os
from PyPDF2 import PdfReader  # For PDF files
import docx  # For DOCX files

# Frontend library
import streamlit as st  # For creating an interactive web interface

# Load the Hugging Face embedding model
embedding_model = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

def read_document(file_path):
    """
    Read document content based on its file format.
    Supports: .txt, .pdf, .docx
    """
    ext = os.path.splitext(file_path)[-1].lower()
    content = ""
    if ext == ".txt":
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
    elif ext == ".pdf":
        reader = PdfReader(file_path)
        for page in reader.pages:
            content += page.extract_text() or ""
    elif ext == ".docx":
        doc = docx.Document(file_path)
        content = " ".join([para.text for para in doc.paragraphs])
    else:
        raise ValueError("Unsupported file format. Supported formats: .txt, .pdf, .docx")
    return content

def split_document(document_text):
    """
    Split the document into chunks and return them.
    """
    text_splitter = CharacterTextSplitter(chunk_size=200, chunk_overlap=0)
    chunks = text_splitter.split_text(document_text)
    return [Document(page_content=chunk) for chunk in chunks]

def search_query(query_text, documents):
    """
    Perform semantic search for a query in the given documents.
    Returns only the top result.
    """
    query_embedding = embedding_model.embed_query(query_text)
    doc_embeddings = embedding_model.embed_documents([doc.page_content for doc in documents])
    similarities = [(doc, sum(q * d for q, d in zip(query_embedding, doc_emb)))
                    for doc, doc_emb in zip(documents, doc_embeddings)]
    similarities.sort(key=lambda x: x[1], reverse=True)
    return similarities[0] if similarities else (None, 0)  # Return (None, 0) if no results

# Streamlit Web Interface
def main():
    st.set_page_config(page_title="Hugging Face Semantic Search", page_icon="🔍", layout="wide")
    st.title("📄 Hugging Face Document Semantic Search")
    st.markdown("""
    <style>
    .reportview-container { background-color: #f0f2f6; }
    .stButton>button { width: 100%; height: 3em; }
    </style>
    """, unsafe_allow_html=True)

    # Initialize session state variables
    if "documents" not in st.session_state:
        st.session_state.documents = []
    if "search_result" not in st.session_state:
        st.session_state.search_result = None

    uploaded_file = st.file_uploader("Upload Document (txt, pdf, docx):", type=["txt", "pdf", "docx"])

    if uploaded_file is not None:
        try:
            with open(uploaded_file.name, "wb") as f:
                f.write(uploaded_file.getbuffer())
            document_text = read_document(uploaded_file.name)
            st.session_state.documents = split_document(document_text)
            st.success("✅ Document successfully processed.")
        except Exception as e:
            st.error(f"⚠️ Error processing file: {e}")

    query_text = st.text_input("Enter your query:", placeholder="Type your semantic search query here...")

    if st.button("🔍 Search"):
        if query_text and st.session_state.documents:
            doc, score = search_query(query_text, st.session_state.documents)
            if doc:
                st.session_state.search_result = (doc.page_content, score)
            else:
                st.session_state.search_result = None

    # Display the search result
    if st.session_state.search_result:
        st.subheader("🔎 Search Result:")
        content, score = st.session_state.search_result
        st.markdown(f"**{content}**")
        st.markdown(f"\n*Similarity Score:* {score:.4f}")
    else:
        st.info("No matching results found.")

if __name__ == "__main__":
    main()